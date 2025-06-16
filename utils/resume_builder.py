from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from io import BytesIO
import tempfile
import traceback
import re

class ValidationError(Exception):
    """Custom exception for validation errors"""
    pass

class ResumeBuilder:
    def __init__(self):
        self.templates = {
            "Modern": self.build_modern_template,
            "Professional": self.build_professional_template,
            "Minimal": self.build_minimal_template,
            "Creative": self.build_creative_template,
            "Executive": self.build_executive_template,
            "Tech": self.build_tech_template
        }
        
    def validate_data(self, data):
        """Enhanced validation with better error handling and more comprehensive checks"""
        errors = []
        warnings = []
        
        # Check if data is provided
        if not data or not isinstance(data, dict):
            raise ValidationError("No data provided for resume generation")
        
        # Validate personal information (required)
        personal_info = data.get('personal_info', {})
        if not personal_info:
            errors.append("Personal information section is required")
        else:
            # Required personal fields
            required_personal_fields = {
                'full_name': 'Full Name'
            }
            
            for field, display_name in required_personal_fields.items():
                value = personal_info.get(field)
                if not value or not str(value).strip():
                    errors.append(f"Personal info: {display_name} is required")
                elif len(str(value).strip()) < 2:
                    errors.append(f"Personal info: {display_name} must be at least 2 characters long")
            
            # Recommended personal fields
            recommended_personal_fields = {
                'email': 'Email address',
                'phone': 'Phone number',
                'location': 'Location'
            }
            
            for field, display_name in recommended_personal_fields.items():
                if not personal_info.get(field) or not str(personal_info[field]).strip():
                    warnings.append(f"Personal info: {display_name} is recommended for better visibility")
            
            # Enhanced email validation
            if personal_info.get('email'):
                email = str(personal_info['email']).strip()
                email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
                if not re.match(email_pattern, email):
                    errors.append("Personal info: Email format is invalid (example: user@domain.com)")
                elif len(email) > 100:
                    warnings.append("Personal info: Email address seems unusually long")
            
            # Enhanced phone validation
            if personal_info.get('phone'):
                phone = str(personal_info['phone']).strip()
                phone_clean = re.sub(r'[^\d+()-\\s]', '', phone)
                if len(phone_clean) < 7:
                    warnings.append("Personal info: Phone number seems too short")
                elif len(phone_clean) > 20:
                    warnings.append("Personal info: Phone number seems too long")
            
            # URL validation for LinkedIn and Portfolio
            for url_field in ['linkedin', 'portfolio', 'github']:
                if personal_info.get(url_field):
                    url = str(personal_info[url_field]).strip()
                    if url and not (url.startswith('http://') or url.startswith('https://') or url.startswith('www.')):
                        warnings.append(f"Personal info: {url_field.title()} should include 'https://' or 'www.'")
        
        # Validate template selection
        available_templates = list(self.templates.keys())
        if not data.get('template'):
            errors.append("Template selection is required")
        elif data['template'] not in available_templates:
            errors.append(f"Invalid template '{data['template']}'. Available templates: {', '.join(available_templates)}")
        
        # Enhanced content validation
        has_experience = data.get('experience') and len(data['experience']) > 0
        has_education = data.get('education') and len(data['education']) > 0
        has_projects = data.get('projects') and len(data['projects']) > 0
        
        if not (has_experience or has_education or has_projects):
            errors.append("At least one content section is required: Experience, Education, or Projects")
        
        # Enhanced experience validation
        if data.get('experience'):
            if not isinstance(data['experience'], list):
                errors.append("Experience section must be a list of experience items")
            else:
                for i, exp in enumerate(data['experience']):
                    if not isinstance(exp, dict):
                        errors.append(f"Experience item {i+1}: Must be a valid object/dictionary")
                        continue
                    
                    required_exp_fields = {
                        'position': 'Job Position',
                        'company': 'Company Name',
                        'start_date': 'Start Date',
                        'end_date': 'End Date'
                    }
                    
                    for field, display_name in required_exp_fields.items():
                        value = exp.get(field)
                        if not value or not str(value).strip():
                            errors.append(f"Experience item {i+1}: {display_name} is required")
                        elif field in ['position', 'company'] and len(str(value).strip()) < 2:
                            errors.append(f"Experience item {i+1}: {display_name} must be at least 2 characters")
                    
                    # Enhanced date validation
                    self._validate_dates(exp, f"Experience item {i+1}", errors, warnings)
                    
                    # Check for description or responsibilities
                    if not exp.get('description') and not exp.get('responsibilities'):
                        warnings.append(f"Experience item {i+1}: Consider adding description or responsibilities")
        
        # Enhanced education validation
        if data.get('education'):
            if not isinstance(data['education'], list):
                errors.append("Education section must be a list of education items")
            else:
                for i, edu in enumerate(data['education']):
                    if not isinstance(edu, dict):
                        errors.append(f"Education item {i+1}: Must be a valid object/dictionary")
                        continue
                    
                    required_edu_fields = {
                        'school': 'School/Institution Name',
                        'degree': 'Degree Type',
                        'field': 'Field of Study',
                        'graduation_date': 'Graduation Date'
                    }
                    
                    for field, display_name in required_edu_fields.items():
                        value = edu.get(field)
                        if not value or not str(value).strip():
                            errors.append(f"Education item {i+1}: {display_name} is required")
                        elif field != 'graduation_date' and len(str(value).strip()) < 2:
                            errors.append(f"Education item {i+1}: {display_name} must be at least 2 characters")
                    
                    # GPA validation
                    if edu.get('gpa'):
                        try:
                            gpa_val = float(edu['gpa'])
                            if gpa_val < 0 or gpa_val > 4.0:
                                warnings.append(f"Education item {i+1}: GPA should typically be between 0.0 and 4.0")
                        except (ValueError, TypeError):
                            warnings.append(f"Education item {i+1}: GPA should be a number")
        
        # Enhanced projects validation
        if data.get('projects'):
            if not isinstance(data['projects'], list):
                errors.append("Projects section must be a list of project items")
            else:
                for i, proj in enumerate(data['projects']):
                    if not isinstance(proj, dict):
                        errors.append(f"Project item {i+1}: Must be a valid object/dictionary")
                        continue
                    
                    if not proj.get('name') or not str(proj['name']).strip():
                        errors.append(f"Project item {i+1}: Project name is required")
                    elif len(str(proj['name']).strip()) < 2:
                        errors.append(f"Project item {i+1}: Project name must be at least 2 characters")
                    
                    if not proj.get('description') and not proj.get('responsibilities'):
                        warnings.append(f"Project item {i+1}: Consider adding description or key features")
        
        # Enhanced skills validation
        if data.get('skills'):
            skills = data['skills']
            if isinstance(skills, dict):
                has_skills = False
                for category in ['technical', 'soft', 'languages', 'tools', 'frameworks', 'databases']:
                    if skills.get(category):
                        skill_items = self._format_list_items(skills[category])
                        if skill_items:
                            has_skills = True
                            break
                
                if not has_skills:
                    warnings.append("Skills section is present but appears empty - consider adding relevant skills")
            else:
                warnings.append("Skills section should be organized by categories (technical, soft, languages, tools)")
        else:
            warnings.append("Skills section is highly recommended to showcase your abilities")
        
        # Professional summary validation
        if data.get('summary'):
            summary = str(data['summary']).strip()
            if len(summary) < 50:
                warnings.append("Professional summary is quite short - consider expanding to 2-3 sentences")
            elif len(summary) > 500:
                warnings.append("Professional summary is quite long - consider condensing to key highlights")
        else:
            warnings.append("Professional summary is recommended to provide a strong opening statement")
        
        # Return comprehensive validation results
        return {
            'is_valid': len(errors) == 0,
            'errors': errors,
            'warnings': warnings,
            'summary': {
                'total_errors': len(errors),
                'total_warnings': len(warnings),
                'sections_present': {
                    'personal_info': bool(personal_info),
                    'summary': bool(data.get('summary')),
                    'experience': bool(data.get('experience')),
                    'education': bool(data.get('education')),
                    'projects': bool(data.get('projects')),
                    'skills': bool(data.get('skills'))
                }
            }
        }
    
    def _validate_dates(self, item, context, errors, warnings):
        """Enhanced date validation helper"""
        for date_field in ['start_date', 'end_date', 'graduation_date']:
            if item.get(date_field):
                date_str = str(item[date_field]).strip()
                
                # Enhanced date patterns
                valid_patterns = [
                    (r'^\d{4}$', '2023'),
                    (r'^\d{1,2}/\d{4}$', '01/2023'),
                    (r'^\d{4}-\d{1,2}$', '2023-01'),
                    (r'^[A-Za-z]{3,9}\\s+\d{4}$', 'January 2023'),
                    (r'^\d{1,2}/\d{1,2}/\d{4}$', '01/15/2023'),
                    (r'^[A-Za-z]{3}\\s+\d{4}$', 'Jan 2023'),
                    (r'^(Present|present|Current|current|Ongoing|ongoing)$', 'Present')
                ]
                
                is_valid = any(re.match(pattern, date_str) for pattern, _ in valid_patterns)
                if not is_valid:
                    example_formats = ', '.join([example for _, example in valid_patterns[:4]])
                    warnings.append(f"{context}: {date_field.replace('_', ' ').title()} format may be incorrect. Try formats like: {example_formats}")
    
    def generate_resume(self, data):
        """Enhanced resume generation with better error handling"""
        try:
            template_name = data.get('template', 'Modern')
            print(f"Starting resume generation with template: {template_name}")
            
            # Validate data first
            validation_result = self.validate_data(data)
            
            if not validation_result['is_valid']:
                error_msg = "Resume generation failed due to validation errors:\n\n"
                error_msg += "ERRORS:\n"
                for i, error in enumerate(validation_result['errors'], 1):
                    error_msg += f"{i}. {error}\n"
                
                if validation_result['warnings']:
                    error_msg += "\nWARNINGS:\n"
                    for i, warning in enumerate(validation_result['warnings'], 1):
                        error_msg += f"{i}. {warning}\n"
                
                error_msg += f"\nValidation Summary: {validation_result['summary']['total_errors']} errors, {validation_result['summary']['total_warnings']} warnings"
                raise ValidationError(error_msg)
            
            # Display warnings if any
            if validation_result['warnings']:
                print(f"Validation completed with {len(validation_result['warnings'])} warnings:")
                for i, warning in enumerate(validation_result['warnings'], 1):
                    print(f"  {i}. {warning}")
            
            # Create document
            doc = Document()
            
            # Apply selected template
            template_function = self.templates.get(template_name.title(), self.build_modern_template)
            doc = template_function(doc, data)
            
            # Save to buffer
            buffer = BytesIO()
            print("Saving document to buffer...")
            doc.save(buffer)
            buffer.seek(0)
            print(f"Resume generated successfully using {template_name} template!")
            
            return buffer
            
        except ValidationError as e:
            print(f"Validation Error: {str(e)}")
            raise
        except Exception as e:
            print(f"Error in generate_resume: {str(e)}")
            print(f"Full traceback: {traceback.format_exc()}")
            raise Exception(f"Resume generation failed: {str(e)}")

    def _format_list_items(self, items):
        """Enhanced helper function to handle various input types"""
        if items is None:
            return []
        
        if isinstance(items, str):
            # Handle both newline and comma-separated items
            if '\n' in items:
                result = [item.strip() for item in items.split('\n') if item.strip()]
            else:
                result = [item.strip() for item in items.split(',') if item.strip()]
            return result
        elif isinstance(items, list):
            return [str(item).strip() for item in items if item and str(item).strip()]
        else:
            return [str(items).strip()] if str(items).strip() else []

    def _safe_get(self, data, key, default=""):
        """Enhanced safe getter with better type handling"""
        try:
            value = data.get(key, default)
            if value is None:
                return default
            return str(value).strip() if value else default
        except (AttributeError, TypeError):
            return default

    def build_modern_template(self, doc, data):
        """Enhanced Modern template with contemporary design"""
        try:
            styles = doc.styles
            
            # Modern color scheme
            primary_color = RGBColor(37, 99, 235)  # Modern blue
            secondary_color = RGBColor(71, 85, 105)  # Slate gray
            accent_color = RGBColor(16, 185, 129)  # Emerald
            text_color = RGBColor(30, 41, 59)  # Dark slate
            
            # Enhanced name style
            name_style = self._get_or_create_style(styles, 'Modern Name Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            name_style.font.size = Pt(28)
            name_style.font.bold = True
            name_style.font.color.rgb = primary_color
            name_style.font.name = 'Segoe UI'
            name_style.paragraph_format.space_after = Pt(4)
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Professional title style
            title_style = self._get_or_create_style(styles, 'Modern Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(14)
            title_style.font.color.rgb = secondary_color
            title_style.font.name = 'Segoe UI'
            title_style.paragraph_format.space_after = Pt(12)

            # Section header style
            section_style = self._get_or_create_style(styles, 'Modern Section Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(16)
            section_style.font.bold = True
            section_style.font.color.rgb = primary_color
            section_style.font.name = 'Segoe UI'
            section_style.paragraph_format.space_before = Pt(20)
            section_style.paragraph_format.space_after = Pt(8)

            # Contact style
            contact_style = self._get_or_create_style(styles, 'Modern Contact Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            contact_style.font.size = Pt(11)
            contact_style.font.name = 'Segoe UI'
            contact_style.font.color.rgb = secondary_color
            contact_style.paragraph_format.space_after = Pt(3)

            # Body text style
            body_style = self._get_or_create_style(styles, 'Modern Body', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.size = Pt(11)
            body_style.font.name = 'Segoe UI'
            body_style.font.color.rgb = text_color
            body_style.paragraph_format.space_after = Pt(6)

            # Header section with modern layout
            name_p = doc.add_paragraph(data['personal_info']['full_name'])
            name_p.style = name_style

            if data['personal_info'].get('title'):
                title_p = doc.add_paragraph(self._safe_get(data['personal_info'], 'title'))
                title_p.style = title_style

            # Modern contact layout
            self._add_modern_contact_info(doc, data['personal_info'], contact_style)

            # Add modern separator line
            separator = doc.add_paragraph('─' * 60)
            separator_style = self._get_or_create_style(styles, 'Modern Separator', WD_STYLE_TYPE.PARAGRAPH)
            separator_style.font.color.rgb = accent_color
            separator_style.paragraph_format.space_after = Pt(16)
            separator.style = separator_style

            # Content sections
            self._add_summary_section(doc, data, section_style, body_style, "PROFESSIONAL SUMMARY")
            self._add_experience_section(doc, data, section_style, body_style, primary_color, accent_color, "modern")
            self._add_projects_section(doc, data, section_style, body_style, primary_color, accent_color, "modern")
            self._add_education_section(doc, data, section_style, body_style, secondary_color, "modern")
            self._add_skills_section(doc, data, section_style, body_style, primary_color, "modern")

            self._set_document_margins(doc, 0.7, 0.7, 0.8, 0.8)
            return doc
            
        except Exception as e:
            print(f"Error in build_modern_template: {str(e)}")
            raise

    def build_executive_template(self, doc, data):
        """New Executive template for senior-level professionals"""
        try:
            styles = doc.styles
            
            # Executive color scheme - sophisticated and professional
            primary_color = RGBColor(17, 24, 39)  # Dark slate
            accent_color = RGBColor(139, 69, 19)  # Saddle brown
            secondary_color = RGBColor(75, 85, 99)  # Gray
            
            # Executive name style - prestigious and bold
            name_style = self._get_or_create_style(styles, 'Executive Name', WD_STYLE_TYPE.PARAGRAPH)
            name_style.font.size = Pt(32)
            name_style.font.bold = True
            name_style.font.color.rgb = primary_color
            name_style.font.name = 'Georgia'
            name_style.paragraph_format.space_after = Pt(6)
            name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Executive title style
            exec_title_style = self._get_or_create_style(styles, 'Executive Title', WD_STYLE_TYPE.PARAGRAPH)
            exec_title_style.font.size = Pt(16)
            exec_title_style.font.color.rgb = accent_color
            exec_title_style.font.name = 'Georgia'
            exec_title_style.paragraph_format.space_after = Pt(16)
            exec_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Executive section style
            section_style = self._get_or_create_style(styles, 'Executive Section', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = accent_color
            section_style.font.name = 'Georgia'
            section_style.paragraph_format.space_before = Pt(24)
            section_style.paragraph_format.space_after = Pt(10)

            # Add executive header with elegant styling
            name_p = doc.add_paragraph(data['personal_info']['full_name'].upper())
            name_p.style = name_style

            if data['personal_info'].get('title'):
                title_p = doc.add_paragraph(self._safe_get(data['personal_info'], 'title'))
                title_p.style = exec_title_style

            # Executive contact information - centered and elegant
            contact_style = self._get_or_create_style(styles, 'Executive Contact', WD_STYLE_TYPE.PARAGRAPH)
            contact_style.font.size = Pt(11)
            contact_style.font.name = 'Georgia'
            contact_style.font.color.rgb = secondary_color
            contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_style.paragraph_format.space_after = Pt(4)

            self._add_executive_contact_info(doc, data['personal_info'], contact_style)

            # Add elegant border
            border_p = doc.add_paragraph('━' * 50)
            border_style = self._get_or_create_style(styles, 'Executive Border', WD_STYLE_TYPE.PARAGRAPH)
            border_style.font.color.rgb = accent_color
            border_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            border_style.paragraph_format.space_after = Pt(20)
            border_p.style = border_style

            # Executive body style
            body_style = self._get_or_create_style(styles, 'Executive Body', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.size = Pt(12)
            body_style.font.name = 'Georgia'
            body_style.font.color.rgb = primary_color
            body_style.paragraph_format.space_after = Pt(8)

            # Content sections with executive styling
            self._add_summary_section(doc, data, section_style, body_style, "EXECUTIVE SUMMARY")
            self._add_experience_section(doc, data, section_style, body_style, primary_color, accent_color, "executive")
            self._add_education_section(doc, data, section_style, body_style, secondary_color, "executive")
            self._add_skills_section(doc, data, section_style, body_style, accent_color, "executive")

            self._set_document_margins(doc, 0.8, 0.8, 1.0, 1.0)
            return doc
            
        except Exception as e:
            print(f"Error in build_executive_template: {str(e)}")
            raise

    def build_tech_template(self, doc, data):
        """New Tech template optimized for technical professionals"""
        try:
            styles = doc.styles
            
            # Tech color scheme - modern and technical
            primary_color = RGBColor(34, 197, 94)  # Green
            secondary_color = RGBColor(71, 85, 105)  # Slate
            accent_color = RGBColor(59, 130, 246)  # Blue
            dark_color = RGBColor(15, 23, 42)  # Dark slate
            
            # Tech name style - clean and modern
            name_style = self._get_or_create_style(styles, 'Tech Name', WD_STYLE_TYPE.PARAGRAPH)
            name_style.font.size = Pt(26)
            name_style.font.bold = True
            name_style.font.color.rgb = dark_color
            name_style.font.name = 'Consolas'
            name_style.paragraph_format.space_after = Pt(4)

            # Tech section style with code-like appearance
            section_style = self._get_or_create_style(styles, 'Tech Section', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = primary_color
            section_style.font.name = 'Consolas'
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(8)

            # Add tech header
            name_p = doc.add_paragraph(f"// {data['personal_info']['full_name']}")
            name_p.style = name_style

            if data['personal_info'].get('title'):
                title_style = self._get_or_create_style(styles, 'Tech Title', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.size = Pt(13)
                title_style.font.color.rgb = secondary_color
                title_style.font.name = 'Consolas'
                title_style.paragraph_format.space_after = Pt(12)
                
                title_p = doc.add_paragraph(f"/* {self._safe_get(data['personal_info'], 'title')} */")
                title_p.style = title_style

            # Tech contact style
            contact_style = self._get_or_create_style(styles, 'Tech Contact', WD_STYLE_TYPE.PARAGRAPH)
            contact_style.font.size = Pt(10)
            contact_style.font.name = 'Consolas'
            contact_style.font.color.rgb = accent_color
            contact_style.paragraph_format.space_after = Pt(3)

            self._add_tech_contact_info(doc, data['personal_info'], contact_style)

            # Tech body style
            body_style = self._get_or_create_style(styles, 'Tech Body', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.size = Pt(11)
            body_style.font.name = 'Calibri'
            body_style.font.color.rgb = dark_color
            body_style.paragraph_format.space_after = Pt(6)

            # Content sections with tech styling
            self._add_summary_section(doc, data, section_style, body_style, "// PROFILE")
            self._add_experience_section(doc, data, section_style, body_style, primary_color, accent_color, "tech")
            self._add_projects_section(doc, data, section_style, body_style, primary_color, accent_color, "tech")
            self._add_skills_section(doc, data, section_style, body_style, primary_color, "tech")
            self._add_education_section(doc, data, section_style, body_style, secondary_color, "tech")

            self._set_document_margins(doc, 0.6, 0.6, 0.8, 0.8)
            return doc
            
        except Exception as e:
            print(f"Error in build_tech_template: {str(e)}")
            raise

    def build_professional_template(self, doc, data):
        """Enhanced Professional template"""
        try:
            styles = doc.styles
            
            # Professional color scheme
            primary_color = RGBColor(0, 51, 102)  # Navy blue
            secondary_color = RGBColor(102, 102, 102)  # Gray
            
            # Professional name style
            name_style = self._get_or_create_style(styles, 'Pro Name Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            name_style.font.size = Pt(26)
            name_style.font.bold = True
            name_style.font.color.rgb = primary_color
            name_style.font.name = 'Times New Roman'
            name_style.paragraph_format.space_after = Pt(6)

            # Section style
            section_style = self._get_or_create_style(styles, 'Pro Section Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(14)
            section_style.font.bold = True
            section_style.font.color.rgb = primary_color
            section_style.font.name = 'Times New Roman'
            section_style.paragraph_format.space_before = Pt(16)
            section_style.paragraph_format.space_after = Pt(8)

            # Add professional header
            name_p = doc.add_paragraph(data['personal_info']['full_name'])
            name_p.style = name_style

            # Contact information
            contact_style = self._get_or_create_style(styles, 'Pro Contact Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            contact_style.font.size = Pt(11)
            contact_style.font.name = 'Times New Roman'
            contact_style.font.color.rgb = secondary_color
            contact_style.paragraph_format.space_after = Pt(4)

            self._add_professional_contact_info(doc, data['personal_info'], contact_style)

            # Professional body style
            # Professional body style
            body_style = self._get_or_create_style(styles, 'Pro Body Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.size = Pt(11)
            body_style.font.name = 'Times New Roman'
            body_style.paragraph_format.space_after = Pt(6)

            # Content sections
            self._add_summary_section(doc, data, section_style, body_style, "PROFESSIONAL SUMMARY")
            self._add_experience_section(doc, data, section_style, body_style, primary_color, secondary_color, "professional")
            self._add_projects_section(doc, data, section_style, body_style, primary_color, secondary_color, "professional")
            self._add_education_section(doc, data, section_style, body_style, secondary_color, "professional")
            self._add_skills_section(doc, data, section_style, body_style, primary_color, "professional")

            self._set_document_margins(doc, 0.75, 0.75, 1.0, 1.0)
            return doc
            
        except Exception as e:
            print(f"Error in build_professional_template: {str(e)}")
            raise

    def build_minimal_template(self, doc, data):
        """Enhanced Minimal template with clean design"""
        try:
            styles = doc.styles
            
            # Minimal color scheme
            primary_color = RGBColor(64, 64, 64)  # Dark gray  
            secondary_color = RGBColor(128, 128, 128)  # Medium gray
            
            # Minimal name style
            name_style = self._get_or_create_style(styles, 'Minimal Name Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            name_style.font.size = Pt(24)
            name_style.font.bold = True
            name_style.font.color.rgb = primary_color
            name_style.font.name = 'Arial'
            name_style.paragraph_format.space_after = Pt(8)

            # Minimal section style
            section_style = self._get_or_create_style(styles, 'Minimal Section Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(12)
            section_style.font.bold = True
            section_style.font.color.rgb = primary_color
            section_style.font.name = 'Arial'
            section_style.paragraph_format.space_before = Pt(14)
            section_style.paragraph_format.space_after = Pt(6)

            # Add minimal header
            name_p = doc.add_paragraph(data['personal_info']['full_name'])
            name_p.style = name_style

            # Minimal contact style
            contact_style = self._get_or_create_style(styles, 'Minimal Contact Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            contact_style.font.size = Pt(10)
            contact_style.font.name = 'Arial'
            contact_style.font.color.rgb = secondary_color
            contact_style.paragraph_format.space_after = Pt(2)

            self._add_minimal_contact_info(doc, data['personal_info'], contact_style)

            # Minimal body style
            body_style = self._get_or_create_style(styles, 'Minimal Body Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.size = Pt(10)
            body_style.font.name = 'Arial'
            body_style.paragraph_format.space_after = Pt(4)

            # Content sections
            self._add_summary_section(doc, data, section_style, body_style, "Summary")
            self._add_experience_section(doc, data, section_style, body_style, primary_color, secondary_color, "minimal")
            self._add_projects_section(doc, data, section_style, body_style, primary_color, secondary_color, "minimal")
            self._add_education_section(doc, data, section_style, body_style, secondary_color, "minimal")
            self._add_skills_section(doc, data, section_style, body_style, primary_color, "minimal")

            self._set_document_margins(doc, 0.5, 0.5, 0.75, 0.75)
            return doc
            
        except Exception as e:
            print(f"Error in build_minimal_template: {str(e)}")
            raise

    def build_creative_template(self, doc, data):
        """Enhanced Creative template with artistic flair"""
        try:
            styles = doc.styles
            
            # Creative color scheme
            primary_color = RGBColor(147, 51, 234)  # Purple
            secondary_color = RGBColor(219, 39, 119)  # Pink
            accent_color = RGBColor(59, 130, 246)  # Blue
            text_color = RGBColor(55, 65, 81)  # Gray
            
            # Creative name style
            name_style = self._get_or_create_style(styles, 'Creative Name Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            name_style.font.size = Pt(30)
            name_style.font.bold = True
            name_style.font.color.rgb = primary_color
            name_style.font.name = 'Trebuchet MS'
            name_style.paragraph_format.space_after = Pt(8)

            # Creative section style
            section_style = self._get_or_create_style(styles, 'Creative Section Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            section_style.font.size = Pt(15)
            section_style.font.bold = True
            section_style.font.color.rgb = secondary_color
            section_style.font.name = 'Trebuchet MS'
            section_style.paragraph_format.space_before = Pt(18)
            section_style.paragraph_format.space_after = Pt(8)

            # Add creative header
            name_p = doc.add_paragraph(data['personal_info']['full_name'])
            name_p.style = name_style

            if data['personal_info'].get('title'):
                title_style = self._get_or_create_style(styles, 'Creative Title', WD_STYLE_TYPE.PARAGRAPH)
                title_style.font.size = Pt(14)
                title_style.font.color.rgb = accent_color
                title_style.font.name = 'Trebuchet MS'
                title_style.paragraph_format.space_after = Pt(12)
                
                title_p = doc.add_paragraph(self._safe_get(data['personal_info'], 'title'))
                title_p.style = title_style

            # Creative contact style
            contact_style = self._get_or_create_style(styles, 'Creative Contact Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            contact_style.font.size = Pt(11)
            contact_style.font.name = 'Trebuchet MS'
            contact_style.font.color.rgb = text_color
            contact_style.paragraph_format.space_after = Pt(4)

            self._add_creative_contact_info(doc, data['personal_info'], contact_style)

            # Add creative divider
            divider_p = doc.add_paragraph('◆ ◇ ◆ ◇ ◆ ◇ ◆ ◇ ◆ ◇ ◆ ◇ ◆ ◇ ◆')
            divider_style = self._get_or_create_style(styles, 'Creative Divider', WD_STYLE_TYPE.PARAGRAPH)
            divider_style.font.color.rgb = primary_color
            divider_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            divider_style.paragraph_format.space_after = Pt(16)
            divider_p.style = divider_style

            # Creative body style
            body_style = self._get_or_create_style(styles, 'Creative Body Enhanced', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.size = Pt(11)
            body_style.font.name = 'Trebuchet MS'
            body_style.font.color.rgb = text_color
            body_style.paragraph_format.space_after = Pt(6)

            # Content sections
            self._add_summary_section(doc, data, section_style, body_style, "✨ CREATIVE PROFILE")
            self._add_experience_section(doc, data, section_style, body_style, primary_color, secondary_color, "creative")
            self._add_projects_section(doc, data, section_style, body_style, primary_color, secondary_color, "creative")
            self._add_education_section(doc, data, section_style, body_style, accent_color, "creative")
            self._add_skills_section(doc, data, section_style, body_style, primary_color, "creative")

            self._set_document_margins(doc, 0.8, 0.8, 0.9, 0.9)
            return doc
            
        except Exception as e:
            print(f"Error in build_creative_template: {str(e)}")
            raise

    # Helper methods for contact information formatting
    def _add_modern_contact_info(self, doc, personal_info, contact_style):
        """Add modern styled contact information"""
        contact_items = []
        
        if personal_info.get('email'):
            contact_items.append(f"✉ {personal_info['email']}")
        if personal_info.get('phone'):
            contact_items.append(f"📞 {personal_info['phone']}")
        if personal_info.get('location'):
            contact_items.append(f"📍 {personal_info['location']}")
        if personal_info.get('linkedin'):
            contact_items.append(f"💼 {personal_info['linkedin']}")
        if personal_info.get('portfolio'):
            contact_items.append(f"🌐 {personal_info['portfolio']}")
        if personal_info.get('github'):
            contact_items.append(f"💻 {personal_info['github']}")

        for i, item in enumerate(contact_items):
            if i < 3:  # First row
                if i == 0:
                    contact_p = doc.add_paragraph(item)
                else:
                    contact_p.add_run(f" • {item}")
                contact_p.style = contact_style
            else:  # Second row
                if i == 3:
                    contact_p2 = doc.add_paragraph(item)
                else:
                    contact_p2.add_run(f" • {item}")
                contact_p2.style = contact_style

    def _add_executive_contact_info(self, doc, personal_info, contact_style):
        """Add executive styled contact information"""
        contact_items = []
        
        if personal_info.get('email'):
            contact_items.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_items.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_items.append(personal_info['location'])
        if personal_info.get('linkedin'):
            contact_items.append(personal_info['linkedin'])

        if contact_items:
            contact_text = " | ".join(contact_items)
            contact_p = doc.add_paragraph(contact_text)
            contact_p.style = contact_style

    def _add_tech_contact_info(self, doc, personal_info, contact_style):
        """Add tech styled contact information"""
        if personal_info.get('email'):
            email_p = doc.add_paragraph(f"email: {personal_info['email']}")
            email_p.style = contact_style
        if personal_info.get('phone'):
            phone_p = doc.add_paragraph(f"phone: {personal_info['phone']}")
            phone_p.style = contact_style
        if personal_info.get('github'):
            github_p = doc.add_paragraph(f"github: {personal_info['github']}")
            github_p.style = contact_style
        if personal_info.get('linkedin'):
            linkedin_p = doc.add_paragraph(f"linkedin: {personal_info['linkedin']}")
            linkedin_p.style = contact_style
        if personal_info.get('portfolio'):
            portfolio_p = doc.add_paragraph(f"portfolio: {personal_info['portfolio']}")
            portfolio_p.style = contact_style

    def _add_professional_contact_info(self, doc, personal_info, contact_style):
        """Add professional styled contact information"""
        contact_line1 = []
        contact_line2 = []
        
        if personal_info.get('email'):
            contact_line1.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_line1.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_line2.append(personal_info['location'])
        if personal_info.get('linkedin'):
            contact_line2.append(personal_info['linkedin'])

        if contact_line1:
            contact_p1 = doc.add_paragraph(" • ".join(contact_line1))
            contact_p1.style = contact_style
        if contact_line2:
            contact_p2 = doc.add_paragraph(" • ".join(contact_line2))
            contact_p2.style = contact_style

    def _add_minimal_contact_info(self, doc, personal_info, contact_style):
        """Add minimal styled contact information"""
        contact_items = []
        
        if personal_info.get('email'):
            contact_items.append(personal_info['email'])
        if personal_info.get('phone'):
            contact_items.append(personal_info['phone'])
        if personal_info.get('location'):
            contact_items.append(personal_info['location'])

        if contact_items:
            contact_text = " | ".join(contact_items)
            contact_p = doc.add_paragraph(contact_text)
            contact_p.style = contact_style

        # Additional links on separate line
        links = []
        if personal_info.get('linkedin'):
            links.append(personal_info['linkedin'])
        if personal_info.get('portfolio'):
            links.append(personal_info['portfolio'])
        
        if links:
            links_p = doc.add_paragraph(" | ".join(links))
            links_p.style = contact_style

    def _add_creative_contact_info(self, doc, personal_info, contact_style):
        """Add creative styled contact information"""
        if personal_info.get('email'):
            email_p = doc.add_paragraph(f"✦ {personal_info['email']}")
            email_p.style = contact_style
        if personal_info.get('phone'):
            phone_p = doc.add_paragraph(f"✦ {personal_info['phone']}")
            phone_p.style = contact_style
        if personal_info.get('location'):
            location_p = doc.add_paragraph(f"✦ {personal_info['location']}")
            location_p.style = contact_style
        if personal_info.get('portfolio'):
            portfolio_p = doc.add_paragraph(f"✦ {personal_info['portfolio']}")
            portfolio_p.style = contact_style

    # Content section methods
    def _add_summary_section(self, doc, data, section_style, body_style, title="SUMMARY"):
        """Add professional summary section"""
        if data.get('summary'):
            section_p = doc.add_paragraph(title)
            section_p.style = section_style
            
            summary_p = doc.add_paragraph(self._safe_get(data, 'summary'))
            summary_p.style = body_style

    def _add_experience_section(self, doc, data, section_style, body_style, primary_color, secondary_color, template_type):
        """Enhanced experience section with template-specific styling"""
        if not data.get('experience'):
            return
            
        # Template-specific section titles
        section_titles = {
            "modern": "PROFESSIONAL EXPERIENCE",
            "executive": "EXECUTIVE EXPERIENCE",
            "tech": "// EXPERIENCE",
            "professional": "WORK EXPERIENCE",
            "minimal": "Experience",
            "creative": "✨ EXPERIENCE"
        }
        
        section_p = doc.add_paragraph(section_titles.get(template_type, "EXPERIENCE"))
        section_p.style = section_style
        
        for exp in data['experience']:
            # Position and company
            if template_type == "tech":
                header = f"{self._safe_get(exp, 'position')} @ {self._safe_get(exp, 'company')}"
            else:
                header = f"{self._safe_get(exp, 'position')} - {self._safe_get(exp, 'company')}"
            
            position_p = doc.add_paragraph()
            position_run = position_p.add_run(header)
            position_run.bold = True
            position_run.font.color.rgb = primary_color
            position_p.style = body_style
            
            # Date range
            start_date = self._safe_get(exp, 'start_date')
            end_date = self._safe_get(exp, 'end_date')
            if start_date and end_date:
                if template_type == "tech":
                    date_text = f"// {start_date} - {end_date}"
                else:
                    date_text = f"{start_date} - {end_date}"
                    
                date_p = doc.add_paragraph(date_text)
                date_run = date_p.runs[0]
                date_run.font.color.rgb = secondary_color
                date_run.italic = True
                date_p.style = body_style
            
            # Location
            if exp.get('location'):
                location_p = doc.add_paragraph(self._safe_get(exp, 'location'))
                location_run = location_p.runs[0]
                location_run.font.color.rgb = secondary_color
                location_p.style = body_style
            
            # Description
            if exp.get('description'):
                desc_p = doc.add_paragraph(self._safe_get(exp, 'description'))
                desc_p.style = body_style
            
            # Responsibilities
            if exp.get('responsibilities'):
                resp_items = self._format_list_items(exp['responsibilities'])
                for item in resp_items:
                    if template_type == "tech":
                        bullet = "→"
                    elif template_type == "creative":
                        bullet = "◦"
                    else:
                        bullet = "•"
                    
                    resp_p = doc.add_paragraph(f"{bullet} {item}")
                    resp_p.style = body_style
                    resp_p.paragraph_format.left_indent = Inches(0.25)
            
            # Add spacing between entries
            doc.add_paragraph()

    def _add_projects_section(self, doc, data, section_style, body_style, primary_color, secondary_color, template_type):
        """Enhanced projects section"""
        if not data.get('projects'):
            return
            
        section_titles = {
            "modern": "KEY PROJECTS",
            "tech": "// PROJECTS",
            "creative": "✨ PROJECTS",
            "executive": "STRATEGIC PROJECTS",
            "professional": "NOTABLE PROJECTS",
            "minimal": "Projects"
        }
        
        section_p = doc.add_paragraph(section_titles.get(template_type, "PROJECTS"))
        section_p.style = section_style
        
        for proj in data['projects']:
            # Project name
            name_p = doc.add_paragraph()
            name_run = name_p.add_run(self._safe_get(proj, 'name'))
            name_run.bold = True
            name_run.font.color.rgb = primary_color
            name_p.style = body_style
            
            # Technologies/Tools
            if proj.get('technologies'):
                tech_items = self._format_list_items(proj['technologies'])
                if tech_items:
                    if template_type == "tech":
                        tech_text = f"Stack: {', '.join(tech_items)}"
                    else:
                        tech_text = f"Technologies: {', '.join(tech_items)}"
                    
                    tech_p = doc.add_paragraph(tech_text)
                    tech_run = tech_p.runs[0]
                    tech_run.font.color.rgb = secondary_color
                    tech_run.italic = True
                    tech_p.style = body_style
            
            # Project description
            if proj.get('description'):
                desc_p = doc.add_paragraph(self._safe_get(proj, 'description'))
                desc_p.style = body_style
            
            # Key features/achievements
            if proj.get('features'):
                feature_items = self._format_list_items(proj['features'])
                for item in feature_items:
                    if template_type == "tech":
                        bullet = "▸"
                    elif template_type == "creative":
                        bullet = "◦"
                    else:
                        bullet = "•"
                    
                    feature_p = doc.add_paragraph(f"{bullet} {item}")
                    feature_p.style = body_style
                    feature_p.paragraph_format.left_indent = Inches(0.25)
            
            # Project URL
            if proj.get('url'):
                url_p = doc.add_paragraph(f"Link: {self._safe_get(proj, 'url')}")
                url_run = url_p.runs[0]
                url_run.font.color.rgb = secondary_color
                url_p.style = body_style
            
            doc.add_paragraph()

    def _add_education_section(self, doc, data, section_style, body_style, color, template_type):
        """Enhanced education section"""
        if not data.get('education'):
            return
            
        section_titles = {
            "modern": "EDUCATION",
            "tech": "// EDUCATION",
            "creative": "✨ EDUCATION",
            "executive": "EDUCATION",
            "professional": "EDUCATION",
            "minimal": "Education"
        }
        
        section_p = doc.add_paragraph(section_titles.get(template_type, "EDUCATION"))
        section_p.style = section_style
        
        for edu in data['education']:
            # Degree and field
            degree_text = f"{self._safe_get(edu, 'degree')} in {self._safe_get(edu, 'field')}"
            degree_p = doc.add_paragraph()
            degree_run = degree_p.add_run(degree_text)
            degree_run.bold = True
            degree_run.font.color.rgb = color
            degree_p.style = body_style
            
            # School
            school_p = doc.add_paragraph(self._safe_get(edu, 'school'))
            school_p.style = body_style
            
            # Graduation date and GPA
            grad_info = []
            if edu.get('graduation_date'):
                grad_info.append(f"Graduated: {edu['graduation_date']}")
            if edu.get('gpa'):
                grad_info.append(f"GPA: {edu['gpa']}")
            
            if grad_info:
                grad_p = doc.add_paragraph(" | ".join(grad_info))
                grad_run = grad_p.runs[0]
                grad_run.font.color.rgb = color
                grad_run.italic = True
                grad_p.style = body_style
            
            # Honors/Awards
            if edu.get('honors'):
                honors_items = self._format_list_items(edu['honors'])
                if honors_items:
                    honors_text = f"Honors: {', '.join(honors_items)}"
                    honors_p = doc.add_paragraph(honors_text)
                    honors_p.style = body_style
            
            doc.add_paragraph()

    def _add_skills_section(self, doc, data, section_style, body_style, color, template_type):
        """Enhanced skills section with categorization"""
        if not data.get('skills'):
            return
            
        section_titles = {
            "modern": "CORE COMPETENCIES",
            "tech": "// SKILLS",
            "creative": "✨ SKILLS & TALENTS",
            "executive": "CORE COMPETENCIES",
            "professional": "SKILLS & EXPERTISE",
            "minimal": "Skills"
        }
        
        section_p = doc.add_paragraph(section_titles.get(template_type, "SKILLS"))
        section_p.style = section_style
        
        skills = data['skills']
        
        # Skill categories mapping
        skill_categories = {
            'technical': 'Technical Skills',
            'programming': 'Programming Languages',
            'languages': 'Languages',
            'soft': 'Soft Skills',
            'tools': 'Tools & Platforms',
            'frameworks': 'Frameworks & Libraries',
            'databases': 'Databases',
            'certifications': 'Certifications'
        }
        
        if isinstance(skills, dict):
            for category, items in skills.items():
                if items:
                    skill_items = self._format_list_items(items)
                    if skill_items:
                        # Category header
                        category_name = skill_categories.get(category, category.title())
                        if template_type == "tech":
                            category_text = f"{category_name}:"
                        else:
                            category_text = f"{category_name}:"
                        
                        cat_p = doc.add_paragraph()
                        cat_run = cat_p.add_run(category_text)
                        cat_run.bold = True
                        cat_run.font.color.rgb = color
                        cat_p.style = body_style
                        
                        # Skills list
                        skills_text = ", ".join(skill_items)
                        skills_p = doc.add_paragraph(skills_text)
                        skills_p.style = body_style
                        skills_p.paragraph_format.left_indent = Inches(0.25)
        else:
            # Handle simple skills list
            skill_items = self._format_list_items(skills)
            if skill_items:
                skills_text = ", ".join(skill_items)
                skills_p = doc.add_paragraph(skills_text)
                skills_p.style = body_style

    # Utility methods
    def _get_or_create_style(self, styles, style_name, style_type):
        """Get existing style or create new one"""
        try:
            return styles[style_name]
        except KeyError:
            return styles.add_style(style_name, style_type)

    def _set_document_margins(self, doc, top, bottom, left, right):
        """Set document margins"""
        try:
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(top)
                section.bottom_margin = Inches(bottom)
                section.left_margin = Inches(left)
                section.right_margin = Inches(right)
        except Exception as e:
            print(f"Warning: Could not set margins: {str(e)}")

# Example usage and testing
if __name__ == "__main__":
    # Sample data for testing
    sample_data = {
        "template": "Modern",
        "personal_info": {
            "full_name": "John Doe",
            "title": "Senior Software Engineer",
            "email": "john.doe@email.com",
            "phone": "(555) 123-4567",
            "location": "San Francisco, CA",
            "linkedin": "https://linkedin.com/in/johndoe",
            "github": "https://github.com/johndoe",
            "portfolio": "https://johndoe.dev"
        },
        "summary": "Experienced software engineer with 8+ years of expertise in full-stack development, cloud architecture, and team leadership. Proven track record of delivering scalable solutions and mentoring development teams.",
        "experience": [
            {
                "position": "Senior Software Engineer",
                "company": "Tech Corp",
                "start_date": "January 2020",
                "end_date": "Present",
                "location": "San Francisco, CA",
                "description": "Lead development of microservices architecture serving 1M+ users daily.",
                "responsibilities": [
                    "Architected and implemented cloud-native solutions using AWS and Kubernetes",
                    "Led a team of 5 developers in agile development practices",
                    "Reduced system latency by 40% through optimization and caching strategies"
                ]
            }
        ],
        "projects": [
            {
                "name": "E-commerce Platform",
                "description": "Full-stack e-commerce solution with real-time inventory management",
                "technologies": ["React", "Node.js", "PostgreSQL", "Redis"],
                "features": [
                    "Real-time inventory tracking",
                    "Payment gateway integration",
                    "Admin dashboard with analytics"
                ],
                "url": "https://github.com/johndoe/ecommerce"
            }
        ],
        "education": [
            {
                "degree": "Bachelor of Science",
                "field": "Computer Science",
                "school": "University of California, Berkeley",
                "graduation_date": "May 2015",
                "gpa": "3.7",
                "honors": ["Dean's List", "Magna Cum Laude"]
            }
        ],
        "skills": {
            "technical": ["JavaScript", "Python", "Java", "React", "Node.js"],
            "tools": ["Docker", "Kubernetes", "AWS", "Git", "Jenkins"],
            "databases": ["PostgreSQL", "MongoDB", "Redis"],
            "soft": ["Leadership", "Problem Solving", "Communication"]
        }
    }
    
    # Test the resume builder
    builder = ResumeBuilder()
    
    try:
        # Validate data
        validation_result = builder.validate_data(sample_data)
        print("Validation Result:")
        print(f"Valid: {validation_result['is_valid']}")
        print(f"Errors: {len(validation_result['errors'])}")
        print(f"Warnings: {len(validation_result['warnings'])}")
        
        if validation_result['warnings']:
            print("\nWarnings:")
            for warning in validation_result['warnings']:
                print(f"  - {warning}")
        
        # Generate resume if valid
        if validation_result['is_valid']:
            buffer = builder.generate_resume(sample_data)
            print(f"\nResume generated successfully! Buffer size: {len(buffer.getvalue())} bytes")
        
    except Exception as e:
        print(f"Error: {str(e)}")